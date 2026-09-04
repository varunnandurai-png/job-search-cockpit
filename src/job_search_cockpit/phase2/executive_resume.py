"""Full Executive Resume Generation from Approved Career Vault Facts.

Produces polished, ATS-compliant, recruiter-ready resumes in both DOCX and PDF formats
aligned with Varun Nanduri's master profile and Southwest visual styling.
"""

from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as WordParagraph
from PIL import Image as PillowImage
from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.enums import TA_CENTER, TA_LEFT  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    HRFlowable,
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

_NAVY = RGBColor(10, 45, 80)
_GOLD = RGBColor(176, 133, 35)
_DARK_GRAY = RGBColor(50, 50, 50)
_WHITE = RGBColor(255, 255, 255)

_PDF_NAVY = colors.HexColor("#0A2D50")
_PDF_GOLD = colors.HexColor("#B08523")
_PDF_CHARCOAL = colors.HexColor("#222222")
_PDF_MUTED = colors.HexColor("#555555")


@dataclass(frozen=True, slots=True)
class ExperienceRole:
    company: str
    title: str
    dates: str
    location: str
    bullets: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class ExecutiveResumeData:
    name: str = "NANDURI VARUN"
    title: str = "Senior Product Manager | FinTech, APIs & Digital Platforms"
    contact_info: str = "Hyderabad, India  •  varun.nanduri495@gmail.com  •  linkedin.com/in/nanduri-varun"
    summary: str = (
        "Senior Product Manager with 11 years of experience leading cross-functional teams across "
        "FinTech, digital lending, payments, and omni-channel commerce. Proven track record scaling "
        "platforms to $250M+ annual GMV, launching automated lead-generation funnels producing 500K+ leads, "
        "lifting customer conversion from 2–3% to 5–7%, and driving multi-year API/platform modernizations "
        "in regulated, high-scale global enterprises."
    )
    competency_groups: tuple[tuple[str, str], ...] = (
        ("Product Leadership & Strategy", "Product Vision & Roadmaps, Backlog Ownership, PRD/BRD, MVP Definition, Go-To-Market, 0-to-1 Delivery, Agile/Scrum/SAFe, Stakeholder Alignment"),
        ("Platform & Integrations", "Third-party REST APIs (JSON/XML), Platform Health & SLA Monitoring, Automated Fault Detection, Real-time Transactional Guardrails, Anti-Fraud"),
        ("Data & Telemetry", "Funnel & Conversion Analytics, Behavioral Telemetry, A/B Experimentation, SQL, Python, Tableau, Looker Executive Dashboards"),
        ("Domain Expertise", "Digital Lending & Mortgage Journeys, Property Search, Omni-Channel Commerce, Last-Mile Delivery, Driver Dispatching, Subscriptions"),
    )
    roles: tuple[ExperienceRole, ...] = (
        ExperienceRole(
            company="JPMorganChase & Co.",
            title="Senior Product Associate / Product Owner",
            dates="June 2024 – Present",
            location="Hyderabad, India",
            bullets=(
                "Own product vision, roadmap, and team backlog for Chase My Home, the digital entry point for mortgage discovery, property search, and loan affordability across customer experience and backend platform.",
                "Improved loan-application submission conversion from 2–3% to 5–7% by overhauling the Property Search journey and restructuring third-party REST API integrations with JSON/XML data exchange.",
                "Launched Chase Agent Express with vendor partner Home Story, creating an automated lead-generation intake funnel that produced 500,000+ customer leads and qualified transactions.",
                "Shipped Mortgage and Affordability Calculators integrating soft-pull credit and multiple income sources with compliant, data-validated launch.",
                "Drive backend ATM operations roadmap, automating incident ticketing and fault detection while owning platform-health metrics, SLA monitoring, and production support for mission-critical banking systems.",
                "Leverage Generative AI tools (ChatGPT, Microsoft Copilot) for discovery synthesis, PRD/user-story drafting, and defining evaluation criteria for AI-assisted features.",
            ),
        ),
        ExperienceRole(
            company="Walmart Global Tech India",
            title="Senior Product Manager / Product Lead — Online Pickup & Delivery",
            dates="July 2021 – June 2024",
            location="Bengaluru, India",
            bullets=(
                "Owned product strategy, roadmap, and backlog for core customer and fulfillment capabilities in Walmart's omni-channel commerce ecosystem.",
                "Managed platform features supporting $250M+ annual GMV across multiple Scrum teams.",
                "Led multi-year CINE 2.0 modernization of Walmart's last-mile delivery platform, improving driver dispatching and delivery workflows.",
                "Improved global Subscriptions experience and automated Order Substitutions logic, reducing delivery friction and supporting customer retention at checkout.",
                "Partnered with risk and anti-fraud teams to productize real-time transactional guardrails and escalation rules that reduced platform leakage and vulnerabilities.",
                "Collaborated with data engineering to build scalable data models, behavioral telemetry, and executive dashboards in Tableau and Looker to guide roadmap decisions.",
                "Led cross-functional Scrum teams, mentored Product Owners and analysts, ran iteration planning/reviews/demos, and shipped with high velocity and quality.",
            ),
        ),
        ExperienceRole(
            company="Ness Digital Engineering",
            title="Senior Business Analyst / Proxy Product Owner",
            dates="August 2020 – April 2021",
            location="Bengaluru, India",
            bullets=(
                "Acted as proxy Product Owner, owning team backlog, sprint deliverables, requirements discovery, user stories, acceptance criteria, and stakeholder alignment across the full Agile cycle.",
                "Led requirement gathering and discovery with business stakeholders, translating needs into prioritized user stories and technical specifications.",
                "Built SQL and Tableau dashboards for enterprise clients, enabling data-driven operational decisions.",
            ),
        ),
        ExperienceRole(
            company="TAIGER / Advantage One",
            title="NLP Engineer / Product Analytics",
            dates="February 2020 – August 2020",
            location="Hyderabad, India",
            bullets=(
                "Applied Python and SQL analytics to customer processing funnels, surfacing bottlenecks and presenting data-driven recommendations to leadership.",
                "Used descriptive analytics to identify throughput issues and optimize operational processing workflows.",
            ),
        ),
        ExperienceRole(
            company="Deloitte Touche Tohmatsu LLC",
            title="Consultant / Analyst — Tax and Automation",
            dates="February 2017 – January 2020",
            location="Hyderabad, India",
            bullets=(
                "Ideated and built a secure financial-document upload prototype feeding extracted client data into enterprise tax software, reducing manual data entry and improving workflow efficiency.",
                "Cut errors from missed tax benefits through comprehensive review of internal preparation documents, client interviews, and liaison with revenue authorities including IRS and CRA.",
            ),
        ),
        ExperienceRole(
            company="Four Clover Realty Private Limited",
            title="Senior Business Analyst / Business Analyst",
            dates="November 2014 – October 2016",
            location="Bengaluru, India",
            bullets=(
                "Analyzed customer demand trends and market data to inform product positioning and pricing strategies for commercial and residential developments.",
                "Built financial models and business performance reports for executive leadership.",
            ),
        ),
    )
    education: tuple[str, ...] = (
        "Indian School of Business (ISB) — Professional Certificate in Product Management (2023 – 2024)",
        "Alliance School of Business — MBA, Finance (2013 – 2015)",
        "Arjun College of Technology — B.Tech, Electrical, Electronics & Communications Engineering (2009 – 2013)",
    )
    certifications: tuple[str, ...] = (
        "Professional Scrum Product Owner I (PSPO I) — Scrum.org",
        "Understanding and Visualizing Data with Python — University of Michigan",
        "Professional Certificate in Product Management — Indian School of Business",
    )


def generate_executive_docx(
    data: ExecutiveResumeData,
    output_path: Path,
    target_company: str = "",
    target_role: str = "",
    headshot_path: Path | None = None,
) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Styles
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    # 1. HEADER BAND
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_before = Pt(4)
    header.paragraph_format.space_after = Pt(2)
    
    name_run = header.add_run(data.name)
    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = _NAVY

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(3)
    display_title = data.title
    if target_role:
        display_title = f"{target_role} | FinTech, APIs & Digital Platforms"
    title_run = title_p.add_run(display_title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = _GOLD

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(10)
    c_run = contact_p.add_run(data.contact_info)
    c_run.font.size = Pt(9.5)
    c_run.font.color.rgb = _DARK_GRAY

    # Helper for Section Headings
    def add_section_heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = _NAVY
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "B08523")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 2. EXECUTIVE SUMMARY
    add_section_heading("Executive Summary")
    sum_p = doc.add_paragraph()
    sum_p.paragraph_format.space_after = Pt(6)
    sum_p.paragraph_format.line_spacing = 1.15
    sum_run = sum_p.add_run(data.summary)
    sum_run.font.size = Pt(9.5)

    # 3. CORE COMPETENCIES
    add_section_heading("Core Competencies")
    for group, skills in data.competency_groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        g_run = p.add_run(f"• {group}: ")
        g_run.bold = True
        g_run.font.size = Pt(9.5)
        s_run = p.add_run(skills)
        s_run.font.size = Pt(9.5)

    # 4. PROFESSIONAL EXPERIENCE
    add_section_heading("Professional Experience")
    for role in data.roles:
        # Role Header
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(7)
        rp.paragraph_format.space_after = Pt(2)
        
        comp_run = rp.add_run(role.company)
        comp_run.bold = True
        comp_run.font.size = Pt(10.5)
        comp_run.font.color.rgb = _NAVY

        title_run = rp.add_run(f" — {role.title}")
        title_run.bold = True
        title_run.font.size = Pt(10)

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(3)
        m_run = meta_p.add_run(f"{role.dates}  |  {role.location}")
        m_run.italic = True
        m_run.font.size = Pt(9)
        m_run.font.color.rgb = _DARK_GRAY

        # Bullets
        for bullet in role.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2.5)
            bp.paragraph_format.line_spacing = 1.15
            brun = bp.add_run(bullet)
            brun.font.size = Pt(9.5)

    # 5. EDUCATION & CERTIFICATIONS
    add_section_heading("Education & Certifications")
    for edu in data.education:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {edu}")
        r.font.size = Pt(9.5)
    for cert in data.certifications:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {cert}")
        r.font.size = Pt(9.5)

    doc.save(str(output_path))


def generate_executive_pdf(
    data: ExecutiveResumeData,
    output_path: Path,
    target_company: str = "",
    target_role: str = "",
    headshot_path: Path | None = None,
) -> None:
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "ExecName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        textColor=_PDF_NAVY,
    )
    title_style = ParagraphStyle(
        "ExecTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        alignment=TA_CENTER,
        textColor=_PDF_GOLD,
        spaceAfter=3,
    )
    contact_style = ParagraphStyle(
        "ExecContact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        textColor=_PDF_MUTED,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "ExecHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=_PDF_NAVY,
        spaceBefore=7,
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "ExecBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=_PDF_CHARCOAL,
        spaceAfter=4,
    )
    role_header_style = ParagraphStyle(
        "ExecRoleHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=_PDF_NAVY,
        spaceBefore=5,
        spaceAfter=1,
    )
    role_meta_style = ParagraphStyle(
        "ExecRoleMeta",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=10.5,
        textColor=_PDF_MUTED,
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "ExecBullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=_PDF_CHARCOAL,
        spaceAfter=2,
        leftIndent=12,
    )

    story: list[object] = []

    # 1. HEADER
    story.append(Paragraph(escape(data.name), name_style))
    display_title = data.title
    if target_role:
        display_title = f"{target_role} | FinTech, APIs & Digital Platforms"
    story.append(Paragraph(escape(display_title), title_style))
    story.append(Paragraph(escape(data.contact_info), contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=_PDF_GOLD, spaceBefore=2, spaceAfter=6))

    def add_pdf_heading(title: str) -> None:
        story.append(Paragraph(escape(title.upper()), heading_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=_PDF_GOLD, spaceBefore=1, spaceAfter=4))

    # 2. EXECUTIVE SUMMARY
    add_pdf_heading("Executive Summary")
    story.append(Paragraph(escape(data.summary), body_style))

    # 3. CORE COMPETENCIES
    add_pdf_heading("Core Competencies")
    for group, skills in data.competency_groups:
        text = f"<b>• {escape(group)}:</b> {escape(skills)}"
        story.append(Paragraph(text, body_style))

    # 4. PROFESSIONAL EXPERIENCE
    add_pdf_heading("Professional Experience")
    for role in data.roles:
        header_text = f"<b>{escape(role.company)}</b> — {escape(role.title)}"
        meta_text = f"{escape(role.dates)}  |  {escape(role.location)}"
        story.append(Paragraph(header_text, role_header_style))
        story.append(Paragraph(meta_text, role_meta_style))
        for bullet in role.bullets:
            story.append(Paragraph(f"• {escape(bullet)}", bullet_style))

    # 5. EDUCATION & CERTIFICATIONS
    add_pdf_heading("Education & Certifications")
    for edu in data.education:
        story.append(Paragraph(f"• {escape(edu)}", body_style))
    for cert in data.certifications:
        story.append(Paragraph(f"• {escape(cert)}", body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    doc.build(story)
