from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from pypdf import PdfReader

from job_search_cockpit.phase1_contract.snapshots import (
    Phase1ResumeFactProjection,
    Phase1ResumeFactSnapshot,
)
from job_search_cockpit.phase2.document_rendering import (
    LocalResumeRenderer,
    extract_docx_text,
    extract_pdf_text,
)
from job_search_cockpit.phase2.resume_documents import build_canonical_resume_document


def test_renderer_creates_readable_equivalent_docx_and_pdf_from_one_model(tmp_path: Path) -> None:
    headshot_path = tmp_path / "placeholder.png"
    Image.new("RGB", (120, 120), color=(210, 220, 230)).save(headshot_path)
    document = build_canonical_resume_document(
        Phase1ResumeFactProjection(
            requirement_ids=("skills.python",),
            facts=(
                Phase1ResumeFactSnapshot(
                    requirement_id="skills.python",
                    claim_id="sanitized-claim-1",
                    revision_id="sanitized-revision-1",
                    support_assertion_id="sanitized-support-1",
                    safe_wording="Built Python services.",
                    employer_key=None,
                    period_start=None,
                    period_end=None,
                ),
            ),
            profile_fingerprint="a" * 64,
            profile_generation=1,
            readiness_fingerprint="b" * 64,
            readiness_generation=1,
            authority_fingerprint="c" * 64,
            authority_generation=1,
            restore_generation=0,
            fingerprint="d" * 64,
        )
    )

    rendered = LocalResumeRenderer().render(
        document=document,
        output_dir=tmp_path / "output",
        stem="Varun_Resume_Acme",
        headshot_path=headshot_path,
    )

    assert rendered.docx_path.exists()
    assert rendered.pdf_path.exists()
    assert extract_docx_text(rendered.docx_path) == document.plain_text
    assert extract_pdf_text(rendered.pdf_path) == document.plain_text

    rendered_docx = Document(rendered.docx_path)
    assert rendered_docx.paragraphs[0].text == ""
    header_properties = rendered_docx.paragraphs[1]._p.get_or_add_pPr()
    shading = header_properties.find(qn("w:shd"))
    border = header_properties.find(f"{qn('w:pBdr')}/{qn('w:bottom')}")
    assert shading is not None
    assert shading.get(qn("w:fill")) == "0A2D50"
    assert border is not None
    assert border.get(qn("w:color")) == "B08523"
    bullet_properties = rendered_docx.styles["List Bullet"]._element.get_or_add_pPr()
    assert bullet_properties.find(qn("w:contextualSpacing")) is None
    first_entry_spacing = rendered_docx.paragraphs[3]._p.get_or_add_pPr().find(
        qn("w:spacing")
    )
    assert first_entry_spacing is not None
    assert first_entry_spacing.get(qn("w:after")) == "100"
    assert first_entry_spacing.get(qn("w:line")) == "276"


def test_renderer_preserves_literal_safe_wording_in_both_formats(tmp_path: Path) -> None:
    headshot_path = tmp_path / "placeholder.png"
    Image.new("RGB", (120, 120), color=(210, 220, 230)).save(headshot_path)
    document = build_canonical_resume_document(
        Phase1ResumeFactProjection(
            requirement_ids=("skills.safe",),
            facts=(
                Phase1ResumeFactSnapshot(
                    requirement_id="skills.safe",
                    claim_id="sanitized-claim-1",
                    revision_id="sanitized-revision-1",
                    support_assertion_id="sanitized-support-1",
                    safe_wording="Built <safe> & reliable systems.",
                    employer_key=None,
                    period_start=None,
                    period_end=None,
                ),
            ),
            profile_fingerprint="a" * 64,
            profile_generation=1,
            readiness_fingerprint="b" * 64,
            readiness_generation=1,
            authority_fingerprint="c" * 64,
            authority_generation=1,
            restore_generation=0,
            fingerprint="d" * 64,
        )
    )

    rendered = LocalResumeRenderer().render(
        document=document,
        output_dir=tmp_path / "output",
        stem="Varun_Resume_Acme",
        headshot_path=headshot_path,
    )

    assert extract_docx_text(rendered.docx_path) == document.plain_text
    assert extract_pdf_text(rendered.pdf_path) == document.plain_text


def test_renderer_keeps_a_multi_page_document_equivalent_across_formats(
    tmp_path: Path,
) -> None:
    requirement_ids = tuple(f"skills.item-{index:02d}" for index in range(1, 33))
    headshot_path = tmp_path / "placeholder.png"
    Image.new("RGB", (120, 120), color=(210, 220, 230)).save(headshot_path)
    document = build_canonical_resume_document(
        Phase1ResumeFactProjection(
            requirement_ids=requirement_ids,
            facts=tuple(
                Phase1ResumeFactSnapshot(
                    requirement_id=requirement_id,
                    claim_id=f"sanitized-claim-{index:02d}",
                    revision_id=f"sanitized-revision-{index:02d}",
                    support_assertion_id=f"sanitized-support-{index:02d}",
                    safe_wording=(
                        f"Delivered synthetic verified outcome {index:02d} with bounded "
                        "evidence, measurable controls, and safe local processing."
                    ),
                    employer_key=None,
                    period_start=None,
                    period_end=None,
                )
                for index, requirement_id in enumerate(requirement_ids, start=1)
            ),
            profile_fingerprint="a" * 64,
            profile_generation=1,
            readiness_fingerprint="b" * 64,
            readiness_generation=1,
            authority_fingerprint="c" * 64,
            authority_generation=1,
            restore_generation=0,
            fingerprint="d" * 64,
        )
    )

    rendered = LocalResumeRenderer().render(
        document=document,
        output_dir=tmp_path / "output",
        stem="Varun_Resume_Multi_Page_QA",
        headshot_path=headshot_path,
    )

    assert extract_docx_text(rendered.docx_path) == document.plain_text
    assert extract_pdf_text(rendered.pdf_path) == document.plain_text
    assert len(Document(rendered.docx_path).paragraphs) == 35
    assert len(PdfReader(rendered.pdf_path).pages) >= 2
