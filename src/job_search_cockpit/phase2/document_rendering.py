from collections.abc import Iterable
from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as WordParagraph
from PIL import Image as PillowImage
from pypdf import PdfReader
from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.enums import TA_CENTER  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from job_search_cockpit.phase2.resume_documents import CanonicalResumeDocument

_NAVY = RGBColor(10, 45, 80)
_GOLD = RGBColor(176, 133, 35)
_WHITE = RGBColor(255, 255, 255)


@dataclass(frozen=True, slots=True)
class RenderedResumeFiles:
    docx_path: Path
    pdf_path: Path


class LocalResumeRenderer:
    def render(
        self,
        *,
        document: CanonicalResumeDocument,
        output_dir: Path,
        stem: str,
        headshot_path: Path,
    ) -> RenderedResumeFiles:
        output_dir.mkdir(parents=True, exist_ok=False)
        docx_path = output_dir / f"{stem}.docx"
        pdf_path = output_dir / f"{stem}.pdf"
        self._render_docx(document, docx_path, headshot_path)
        self._render_pdf(document, pdf_path, headshot_path)
        return RenderedResumeFiles(docx_path=docx_path, pdf_path=pdf_path)

    @staticmethod
    def _render_docx(
        document: CanonicalResumeDocument, output_path: Path, headshot_path: Path
    ) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.15
        bullet = doc.styles["List Bullet"]
        bullet.font.name = "Arial"
        bullet.font.size = Pt(10.5)
        bullet.paragraph_format.space_after = Pt(5)
        bullet.paragraph_format.line_spacing = 1.15
        bullet_properties = bullet._element.get_or_add_pPr()
        contextual_spacing = bullet_properties.find(qn("w:contextualSpacing"))
        if contextual_spacing is not None:
            bullet_properties.remove(contextual_spacing)

        top_spacer = doc.add_paragraph()
        top_spacer.paragraph_format.space_after = Pt(0)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(10)
        header.add_run().add_picture(str(headshot_path), width=Inches(0.9))
        header.add_run().add_break()
        title_run = header.add_run(document.title)
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = _WHITE
        _style_header_band(header)
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run(document.section_title)
        heading_run.bold = True
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = _NAVY
        for entry in document.entries:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(entry.safe_wording)
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
        doc.save(str(output_path))

    @staticmethod
    def _render_pdf(
        document: CanonicalResumeDocument, output_path: Path, headshot_path: Path
    ) -> None:
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "resume-title",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=colors.white,
            spaceAfter=4,
        )
        heading_style = ParagraphStyle(
            "resume-heading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#0A2D50"),
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "resume-body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=10.5 * 1.15,
            spaceAfter=5,
        )
        photo_width, photo_height = _scaled_photo_size(headshot_path)
        photo = Image(str(headshot_path), width=photo_width, height=photo_height)
        header = Table(
            [[photo], [Paragraph(escape(document.title), title_style)]],
            colWidths=[7.1 * inch],
        )
        header.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0A2D50")),
                    ("LINEBELOW", (0, -1), (-1, -1), 1.5, colors.HexColor("#B08523")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        story = [
            header,
            Spacer(1, 0.16 * inch),
            Paragraph(escape(document.section_title), heading_style),
        ]
        story.extend(
            Paragraph(escape(entry.safe_wording), body_style, bulletText="•")
            for entry in document.entries
        )
        SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            leftMargin=0.7 * inch,
            rightMargin=0.7 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
        ).build(story)


def extract_docx_text(path: Path) -> str:
    return _normalise_lines(paragraph.text for paragraph in Document(str(path)).paragraphs)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return _normalise_pdf_lines(page.extract_text() or "" for page in reader.pages)


def _normalise_lines(parts: Iterable[str]) -> str:
    lines: list[str] = []
    for part in parts:
        for line in str(part).splitlines():
            normalised = " ".join(line.split())
            if normalised and normalised != "━":
                lines.append(normalised.removeprefix("• ").removeprefix("\x7f "))
    return "\n\n".join(lines)


def _normalise_pdf_lines(parts: Iterable[str]) -> str:
    paragraphs: list[str] = []
    in_entries = False
    for part in parts:
        for line in str(part).splitlines():
            normalised = " ".join(line.split())
            if not normalised or normalised == "━":
                continue
            is_entry = normalised.startswith(("• ", "\x7f "))
            normalised = normalised.removeprefix("• ").removeprefix("\x7f ")
            if is_entry:
                paragraphs.append(normalised)
                in_entries = True
            elif in_entries:
                paragraphs[-1] = f"{paragraphs[-1]} {normalised}"
            else:
                paragraphs.append(normalised)
    return "\n\n".join(paragraphs)


def _style_header_band(paragraph: WordParagraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "0A2D50")
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "B08523")
    borders.append(bottom)
    properties.append(borders)


def _scaled_photo_size(path: Path) -> tuple[float, float]:
    with PillowImage.open(path) as image:
        width, height = image.size
    scale = min((0.9 * inch) / width, (0.9 * inch) / height)
    return width * scale, height * scale
